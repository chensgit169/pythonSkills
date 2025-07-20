from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt, RGBColor

# Create new document
doc2 = Document()

# Set page margins
for section in doc2.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add a table for the header bar
table = doc2.add_table(rows=1, cols=2)
table.autofit = False
table.columns[0].width = Inches(5.5)
table.columns[1].width = Inches(1.5)

# Style for the whole row (background color light green)
shading_elm_1 = OxmlElement('w:shd')
shading_elm_1.set(qn('w:fill'), "DDEFE3")  # Light green background
shading_elm_2 = OxmlElement('w:shd')
shading_elm_2.set(qn('w:fill'), "DDEFE3")

# Left cell (Title)
cell1 = table.cell(0, 0)
cell1.text = ""
cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
cell1._element.get_or_add_tcPr().append(shading_elm_1)
p1 = cell1.paragraphs[0]
run1 = p1.add_run("Motivation Letter for ")
run1.font.bold = True
run1.font.size = Pt(20)
run1.font.color.rgb = RGBColor(0, 0, 0)

run2 = p1.add_run("Deutschlandstipendium")
run2.font.bold = True
run2.font.size = Pt(20)
run2.font.color.rgb = RGBColor(0, 128, 0)  # Green
p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Right cell (Logo placeholder)
cell2 = table.cell(0, 1)
cell2.text = ""
cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
cell2._element.get_or_add_tcPr().append(shading_elm_2)
p2 = cell2.paragraphs[0]
run3 = p2.add_run("Cake\nResume")
run3.font.size = Pt(12)
run3.font.bold = True
run3.font.color.rgb = RGBColor(255, 255, 255)
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Add some space after the header
doc2.add_paragraph("")

# Add the body (same content as before)
body = doc2.add_paragraph()
body.style.font.name = 'Times New Roman'
body.style.font.size = Pt(11)

body_text = """
[Your Full Name]
[Street Address]
[City, Postal Code]
[Country]
[Email Address]
[Phone Number]

[Date]

[Recipient's Name or Title]
[University Name]
[Faculty/Department (if applicable)]
[University Address]
[City, Postal Code]
[Country]

Subject: Motivation Letter for Master’s Program in [Program Name]

Dear [Mr./Mrs./Dr. Last Name / Admissions Committee],

My name is [Your Name], and I am writing to express my strong interest in the Master’s program in [Program Name] at [University Name], starting in [Month Year]. I have long aspired to [your long-term goal], and believe that your program—with its focus on [mention a key strength]—will be the ideal foundation for achieving this.

I graduated with a [Your Degree] in [Your Major] from [Your University] in [Year]. During my studies, I worked on [project or thesis topic], which sparked my passion for [specific field or research area]. This experience allowed me to build solid skills in [mention techniques, tools, or theory], reinforcing my readiness for postgraduate study.

Beyond academics, I have engaged in [internship, volunteer work, extracurricular activity], where I developed strong [e.g., teamwork, leadership, time‑management] skills. For example, [describe a brief anecdote: what you did, what you learned, impact achieved].

I am particularly drawn to [University Name] because of [university or program feature]. I am confident that this program will help me grow both personally and academically, and that I can contribute through [mention what you bring to the table].

Thank you for considering my application. I would be honored to have the opportunity to discuss my background and aspirations further, and am available for an interview at your convenience.

Sincerely,

[Your Name]
"""
body.add_run(body_text)

# Save the new document
output_path2 = "./Motivation_Letter_Styled_Header.docx"
doc2.save(output_path2)
